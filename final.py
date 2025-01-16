import pandas as pd
import random
import openpyxl
from openpyxl.drawing.image import Image as ExcelImage
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image as PILImage
import io
import roman
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def load_question_bank(file_name):
    return pd.read_excel(file_name, sheet_name='Question Bank')

def select_random_question(unit_questions, remaining_marks):
    available_questions = unit_questions[unit_questions['Marks'] <= remaining_marks]
    return None if available_questions.empty else available_questions.sample(1)

def select_questions_for_unit(questions_pool, unit_marks):
    selected_questions = []
    total_marks_selected = 0
    retries = 0
    max_retries = 50

    while total_marks_selected < unit_marks and retries < max_retries:
        remaining = unit_marks - total_marks_selected
        selected_question = select_random_question(questions_pool, remaining)

        if selected_question is not None:
            marks = selected_question['Marks'].values[0]
            selected_questions.append(selected_question)
            total_marks_selected += marks
        else:
            retries += 1
            if retries >= max_retries:
                print(f"Backtracking failed after {max_retries} retries for this unit.")
                return None
            selected_questions = []
            total_marks_selected = 0

    return (pd.concat(selected_questions), total_marks_selected) if total_marks_selected == unit_marks else None

def calculate_fitness(total_marks, easy_marks, medium_marks, theory_marks, numerical_marks, 
                      total_marks_target, easy_range, medium_range, theory_percentage, numerical_percentage):
    """
    Calculate a fitness score based on how close the generated paper is to the desired criteria.
    
    Args:
        total_marks (int): Total marks of the generated paper
        easy_marks (int): Marks of easy questions
        medium_marks (int): Marks of medium questions
        theory_marks (int): Marks of theory questions
        numerical_marks (int): Marks of numerical questions
        total_marks_target (int): Target total marks
        easy_range (tuple): Acceptable range for easy marks
        medium_range (tuple): Acceptable range for medium marks
        theory_percentage (int): Target theory percentage
        numerical_percentage (int): Target numerical percentage
    
    Returns:
        dict: Detailed fitness information
    """
    # Calculate actual percentages
    total_marks_deviation = abs(total_marks - total_marks_target) / total_marks_target * 100
    
    # Easy marks fitness
    easy_fitness = 100
    if easy_marks < easy_range[0] or easy_marks > easy_range[1]:
        easy_fitness = max(0, 100 - (min(abs(easy_marks - easy_range[0]), abs(easy_marks - easy_range[1])) / easy_range[1] * 100))
    
    # Medium marks fitness
    medium_fitness = 100
    if medium_marks < medium_range[0] or medium_marks > medium_range[1]:
        medium_fitness = max(0, 100 - (min(abs(medium_marks - medium_range[0]), abs(medium_marks - medium_range[1])) / medium_range[1] * 100))
    
    # Theory and Numerical percentage fitness
    theory_actual = (theory_marks / total_marks * 100)
    numerical_actual = (numerical_marks / total_marks * 100)
    
    theory_fitness = max(0, 100 - abs(theory_actual - theory_percentage))
    numerical_fitness = max(0, 100 - abs(numerical_actual - (100 - theory_percentage)))
    
    # Calculate overall fitness
    fitness_components = [
        (total_marks_deviation, 20),  # Total marks deviation
        (100 - easy_fitness, 20),     # Easy marks deviation
        (100 - medium_fitness, 20),   # Medium marks deviation
        (100 - theory_fitness, 20),   # Theory percentage deviation
        (100 - numerical_fitness, 20) # Numerical percentage deviation
    ]
    
    # Weighted fitness calculation
    overall_fitness = max(0, 100 - sum(dev * weight / 100 for dev, weight in fitness_components))
    
    return {
        'overall_fitness': overall_fitness,
        'details': {
            'total_marks_deviation': total_marks_deviation,
            'easy_fitness': easy_fitness,
            'medium_fitness': medium_fitness,
            'theory_fitness': theory_fitness,
            'numerical_fitness': numerical_fitness,
            'actual_easy_marks': easy_marks,
            'actual_medium_marks': medium_marks,
            'actual_theory_percentage': theory_actual,
            'actual_numerical_percentage': numerical_actual
        }
    }

def generate_question_paper(file_name, unitwise_marks, easy_range, medium_range, theory_percentage, max_retries=50):
    # Load the question bank
    question_bank = load_question_bank(file_name)

    # Load Theory/Numerical percentage criteria
    general_info = pd.read_excel(file_name, sheet_name="Question Paper - General Inform", header=None)
    placeholders = dict(zip(general_info[0], general_info[1]))
    theory_percent = int(theory_percentage) 
    numerical_percent = 100 - theory_percent  

    theory_range = (
        int(theory_percent - 10), 
        int(theory_percent + 10)
    )
    numerical_range = (
        int(numerical_percent - 10), 
        int(numerical_percent + 10)
    )

    total_marks = sum(unitwise_marks.values())
    
    # Track best-fitting papers
    best_papers = []

    # Detailed logging for each attempt
    attempts_log = []

    for retry in range(max_retries):
        final_selected_questions = []
        unitwise_total_marks = {}
        difficulty_marks = {'Easy': 0, 'Medium': 0, 'Hard': 0}
        question_type_counts = {'Theory': 0, 'Numerical': 0}

        # Unit-wise selection
        for unit, unit_total_marks in unitwise_marks.items():
            unit_questions = question_bank[question_bank['Unit_No'] == unit]
            result = select_questions_for_unit(unit_questions, unit_total_marks)

            if result is None:
                print(f"Error: Could not meet the total marks requirement for Unit {unit}.")
                break

            selected_unit_questions, total_marks_selected = result
            final_selected_questions.append(selected_unit_questions)
            unitwise_total_marks[unit] = total_marks_selected

            # Update difficulty and question type counts
            for _, row in selected_unit_questions.iterrows():
                difficulty_marks[row['Diff_Level']] += row['Marks']
                question_type_counts[row['Question_Type']] += row['Marks']

        # If unit selection failed, continue to next iteration
        if len(final_selected_questions) != len(unitwise_marks):
            continue

        # Validate criteria
        easy_marks = difficulty_marks['Easy']
        medium_marks = difficulty_marks['Medium']
        theory_marks = question_type_counts['Theory']
        numerical_marks = question_type_counts['Numerical']

        # Calculate Theory and Numerical percentages
        theory_percentage_actual = (theory_marks / total_marks) * 100
        numerical_percentage_actual = (numerical_marks / total_marks) * 100

        # Calculate fitness
        fitness_result = calculate_fitness(
            total_marks=total_marks,
            easy_marks=easy_marks,
            medium_marks=medium_marks,
            theory_marks=theory_marks,
            numerical_marks=numerical_marks,
            total_marks_target=total_marks,
            easy_range=easy_range,
            medium_range=medium_range,
            theory_percentage=theory_percent,
            numerical_percentage=numerical_percent
        )

        # Log attempt details
        attempt_log = {
            'attempt': retry + 1,
            'difficulty_marks': difficulty_marks,
            'theory_percentage': theory_percentage_actual,
            'numerical_percentage': numerical_percentage_actual,
            'fitness': fitness_result
        }
        attempts_log.append(attempt_log)

        # Print detailed logging for each attempt
        print(f"\n--- Attempt {retry + 1} Details ---")
        print("Difficulty Marks:")
        print(f"  Easy: {difficulty_marks['Easy']}")
        print(f"  Medium: {difficulty_marks['Medium']}")
        print(f"  Hard: {difficulty_marks['Hard']}")
        print("\nPercentages:")
        print(f"  Theory: {theory_percentage_actual:.2f}%")
        print(f"  Numerical: {numerical_percentage_actual:.2f}%")
        print("\nFitness Analysis:")
        print(f"  Overall Fitness: {fitness_result['overall_fitness']:.2f}%")

        # Store the generated paper in best papers list
        combined_questions = pd.concat(final_selected_questions)
        best_papers.append({
            'questions': combined_questions,
            'fitness': fitness_result['overall_fitness'],
            'details': fitness_result['details']
        })

        # If all criteria are met, return this paper
        if (easy_range[0] <= easy_marks <= easy_range[1] and
            medium_range[0] <= medium_marks <= medium_range[1] and
            theory_range[0] <= theory_percentage_actual <= theory_range[1] and
            numerical_range[0] <= numerical_percentage_actual <= numerical_range[1]):
            
            print("\n--- Optimal Paper Found ---")
            print_paper_details(combined_questions, unitwise_total_marks, difficulty_marks, 
                                 theory_marks, numerical_marks, total_marks, retry)
            return combined_questions

    # If max retries reached, select from best papers
    if best_papers:
        # Sort papers by fitness in descending order
        best_papers.sort(key=lambda x: x['fitness'], reverse=True)
        
        # Select top 3 papers
        top_papers = best_papers[:3]
        
        # Print details of top 3 papers
        print("\n--- Top 3 Best Fitting Papers ---")
        for i, paper in enumerate(top_papers, 1):
            print(f"\nPaper {i}:")
            print(f"Fitness Score: {paper['fitness']:.4f}")
            for key, value in paper['details'].items():
                print(f"  {key.replace('_', ' ').title()}: {value:.2f}")
        
        # Randomly select one of the top 3 papers
        selected_paper = random.choice(top_papers)
        
        print("\n--- Randomly Selected Best-Fitting Paper ---")
        print_paper_details(selected_paper['questions'], 
                            {unit: sum(group['Marks']) for unit, group in selected_paper['questions'].groupby('Unit_No')}, 
                            selected_paper['questions'].groupby('Diff_Level')['Marks'].sum(), 
                            selected_paper['questions'][selected_paper['questions']['Question_Type'] == 'Theory']['Marks'].sum(),
                            selected_paper['questions'][selected_paper['questions']['Question_Type'] == 'Numerical']['Marks'].sum(), 
                            total_marks, max_retries)
        
        return selected_paper['questions']
    
    print("Unable to generate a suitable question paper after maximum retries.")
    return None

def print_paper_details(questions, unitwise_total_marks, difficulty_marks, theory_marks, numerical_marks, total_marks, retries):
    """
    Print detailed information about the generated question paper.
    """
    print("\n--- Final Selected Questions ---")
    for _, question in questions.iterrows():
        print(f"Unit: {question['Unit_No']}, Marks: {question['Marks']}, "
              f"Difficulty: {question['Diff_Level']}, Type: {question['Question_Type']}, "
              f"Question: {question['Question']}")

    print("\n--- Total Marks Per Unit ---")
    for unit, marks in unitwise_total_marks.items():
        print(f"Unit {unit}: {marks} Marks")

    print("\n--- Total Marks Per Difficulty Level ---")
    if isinstance(difficulty_marks, dict):
        for difficulty, marks in difficulty_marks.items():
            print(f"{difficulty}: {marks} Marks")
    else:
        print(difficulty_marks)

    print("\n--- Theory/Numerical Distribution ---")
    print(f"Theory: {theory_marks} Marks ({theory_marks / total_marks * 100:.2f}%)")
    print(f"Numerical: {numerical_marks} Marks ({numerical_marks / total_marks * 100:.2f}%)")

    print(f"\nTotal Marks Across All Units: {total_marks} Marks")
    print("\nValid paper generated after {} tries.".format(retries + 1))


def convert_excel_to_word_with_images_and_equations(ws, row_number, paragraph):
    headers = {cell.value: idx for idx, cell in enumerate(ws[1])}
    
    if 'Question' not in headers:
        print("Error: 'Question' column is not found.")
        return

    row = ws[row_number]
    question_text = row[headers['Question']].value if row[headers['Question']].value else ""

    while '{{' in question_text and '}}' in question_text:
        start_idx = question_text.index('{{')
        end_idx = question_text.index('}}') + 2
        placeholder = question_text[start_idx:end_idx]
        
        before_placeholder = question_text[:start_idx]
        if before_placeholder:
            run = paragraph.add_run(before_placeholder)
            run.font.size = Pt(11)  # Set font size for text

        image_col_name = placeholder.strip('{}')
        
        if image_col_name in headers:
            for image in ws._images:
                if (image.anchor._from.row == row_number - 1 and 
                    image.anchor._from.col == headers[image_col_name]):
                    img_path = image.ref
                    pil_image = PILImage.open(img_path)
                    image_stream = io.BytesIO()
                    pil_image.save(image_stream, format='PNG')
                    image_stream.seek(0)
                    run = paragraph.add_run()
                    run.add_picture(image_stream, width=Inches(2.5))  # Adjusted image width
                    
        question_text = question_text[end_idx:]
    
    if question_text:
        run = paragraph.add_run(question_text)
        run.font.size = Pt(11)  # Set font size for remaining text

def roman_to_int(roman):
    """Convert a Roman numeral to an integer."""
    roman_map = {'I': 1, 'V': 5, 'X': 10, 'L': 50, 'C': 100, 'D': 500, 'M': 1000}
    value = 0
    prev_value = 0
    for char in reversed(roman):
        current_value = roman_map.get(char, 0)
        if current_value < prev_value:
            value -= current_value
        else:
            value += current_value
        prev_value = current_value
    return value

def preprocess_placeholders_for_paragraphs(placeholders):
    """Create a modified placeholders dictionary for paragraphs."""
    updated_placeholders = placeholders.copy()
    if 'Semester' in updated_placeholders:
        roman_value = str(updated_placeholders['Semester']).strip()  # Extract the Roman numeral
        int_value = roman_to_int(roman_value)
        updated_placeholders['Semester'] = "Odd" if int_value % 2 != 0 else "Even"
    return updated_placeholders
    
def replace_placeholders_in_template(doc, question_bank,set_number, font_size):
    # Load general information from the Excel sheet
    sheet_name = "Question Paper - General Inform"
    general_info = pd.read_excel(question_bank, sheet_name=sheet_name, header=None)
    placeholders = dict(zip(general_info[0], general_info[1]))
    placeholders.update({'Set': set_number})
    # Create a modified placeholders dictionary for paragraphs
    paragraph_placeholders = preprocess_placeholders_for_paragraphs(placeholders)
    
    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        for key, value in paragraph_placeholders.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(value))
                for run in para.runs:
                    if key in ['Total Time', 'Total Marks']:
                        run.font.size = Pt(13)
                        run.font.bold = True
                        run.font.italic = True
                    elif key == 'General Instructions':
                        run.font.size = Pt(12)
                        run.font.bold = True
                    elif key == 'Semester':
                        run.font.size = Pt(12)
                        run.font.bold = True
                    else:
                        run.font.size = Pt(font_size)
                        run.font.bold = True
                    run.font.name = 'Times New Roman'
    
    # Replace placeholders in tables (use original placeholders)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in placeholders.items():  # Use original placeholders
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, str(value))
                            for run in para.runs:
                                run.font.size = Pt(12)
                                run.font.name = 'Times New Roman'
    
    # Second loop - Replace other placeholders with Excel values
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            for key, value in paragraph_placeholders.items():
                if key != 'Set':  # Skip Set placeholder as it's handled above
                    placeholder = f"{{{{{key}}}}}"  # Creates {{key}} format
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, str(value))
                        for run in para.runs:
                            run.font.size = Pt(12)
                            run.font.name = 'Times New Roman'
                else:
                    print(f"Key not found in placeholders: {key}")
                

def apply_table_styles(table):
    # Apply table-wide styles
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Style the header row
    for cell in table.rows[0].cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(13)

    # Style data rows
    for row in table.rows[1:]:
        for i, cell in enumerate(row.cells):
            # Center align Q.No., CO, and Marks columns
            if i in [0, 2, 3]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Left align Question column
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Set font size
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

def get_column_widths(table):
    """
    Get the widths of columns in a table.
    
    Args:
        table: A docx table object
        
    Returns:
        list: List of column widths in inches
    """
    widths = []
    for cell in table.rows[0].cells:
        if cell.width:
            widths.append(cell.width)
        else:
            # Default width if not specified
            widths.append(Inches(1))
    return widths

# def set_column_widths(table, widths):
#     """
#     Set the widths of columns in a table.
    
#     Args:
#         table: A docx table object
#         widths (list): List of column widths to set
#     """
#     for i, width in enumerate(widths):
#         for row in table.rows:
#             row.cells[i].width = width+1


def set_column_widths(table, widths):
    """
    Set the widths of columns in a table.
    
    Args:
        table: A docx table object
        widths (list): List of column widths in inches
    """
    for i, width in enumerate(widths):
        for row in table.rows:
            cell = row.cells[i]
            # Access the XML element for the cell width
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'dxa')
            tcW.set(qn('w:w'), str(int(width * 1440)))  # Inches to Twips (1 inch = 1460 twips)
            tcPr.append(tcW)

def create_word_document_with_images(final_questions, excel_file, word_file, template_file,set_number):
    wb = openpyxl.load_workbook(excel_file)
    ws = wb['Question Bank']
    doc = Document(template_file)

    # Replace placeholders
    replace_placeholders_in_template(doc, excel_file,set_number, 16)

    # Store the widths from the third table before clearing placeholders
    template_table_widths = None
    if len(doc.tables) >= 4:  # Check if there's a fourth table (index 3)
        template_table_widths = get_column_widths(doc.tables[3])
        print(f"Retrieved template table widths: {template_table_widths}")
    else:
        print("Warning: Third table not found in template. Using default widths.")
        template_table_widths = [Inches(0.6), Inches(4.7), Inches(0.6), Inches(0.6)]

    # Clear {{Question Here}} placeholder
    for para in doc.paragraphs:
        if "{{Question Here}}" in para.text:
            para.text = ""

    # Add section for each unit
    for unit in final_questions['Unit_No'].unique():
        # Add unit header with proper formatting
        unit_header = doc.add_paragraph()
        unit_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        unit_run = unit_header.add_run(f"Q. {unit}: Answer the Following")
        unit_run.font.bold = True
        unit_run.font.size = Pt(12)
        unit_header.paragraph_format.space_after = Pt(12)

        # Create and configure table
        table = doc.add_table(rows=1, cols=4)
        table.allow_autofit = False  # Disable autofit to maintain fixed widths
        
        # Apply the stored widths
        set_column_widths(table, template_table_widths)

        # Add header row
        header_cells = table.rows[0].cells
        headers = ['Q. No.', 'Question', 'CO', 'Marks']
        for i, header in enumerate(headers):
            header_cells[i].text = header

        # Add questions
        subquestion = 0
        unit_questions = final_questions[final_questions['Unit_No'] == unit]
        for idx, question_row in unit_questions.iterrows():
            cells = table.add_row().cells

            # Set Q. No. (A, B, C, etc.)
            cells[0].text = chr(65 + subquestion % 26)
            subquestion += 1

            # Add question content
            question_para = cells[1].paragraphs[0]
            convert_excel_to_word_with_images_and_equations(ws, question_row.name + 2, question_para)

            # Set CO and Marks
            cells[2].text = str(question_row['CO'])
            cells[3].text = str(int(question_row['Marks']))

            # Reapply widths to ensure consistency after adding content
            set_column_widths(table, template_table_widths)

        # Apply table styles
        apply_table_styles(table)
        
        # Add spacing after table
        doc.add_paragraph()

    # Add end of paper
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_run = end_para.add_run("*** End of Paper ***")
    end_run.font.bold = True
    end_run.font.size = Pt(12)
    doc.tables[2]._element.getparent().remove(doc.tables[2]._element)
    
    # Save document
    doc.save(word_file)
    print(f'Question paper has been successfully generated: {word_file}')





def apply_table_styles_master(table):
    # Apply table-wide styles
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Style the header row
    for cell in table.rows[0].cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(13)

    # Style data rows
    for row in table.rows[1:]:
        for i, cell in enumerate(row.cells):
            # Center align Q.No., Diff_Level, Blooms_Level, CO, and Marks columns
            if i in [0, 2, 3, 4, 5]:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Left align Question column
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Set font size
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    # Define widths in inches for the 6 columns
    widths = [
        Inches(0.5),   # Q. No.
        Inches(4.5),   # Question
        Inches(0.58),  # Diff_Level
        Inches(0.58),  # Blooms_Level
        Inches(0.53),  # CO
        Inches(0.53)   # Marks
    ]
    
    # Apply column widths
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

def create_word_document_master_with_images(final_questions, excel_file, word_file, template_file, set_number):
    wb = openpyxl.load_workbook(excel_file)
    ws = wb['Question Bank']
    doc = Document(template_file)

    # Replace placeholders
    replace_placeholders_in_template(doc, excel_file, set_number, 16)

    # Store the widths from the third table before clearing placeholders
    template_table_widths = None
    if len(doc.tables) >= 4:
        template_table_widths = get_column_widths(doc.tables[3])
        print(f"Retrieved template table widths: {template_table_widths}")
    else:
        print("Warning: Third table not found in template. Using default widths.")
        template_table_widths = [Inches(0.5), Inches(4.5), Inches(0.8), Inches(0.8), Inches(0.53), Inches(0.53)]

    # Clear {{Question Here}} placeholder
    for para in doc.paragraphs:
        if "{{Question Here}}" in para.text:
            para.text = ""

    # Add section for each unit
    for unit in final_questions['Unit_No'].unique():
        # Add unit header with proper formatting
        unit_header = doc.add_paragraph()
        unit_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        unit_run = unit_header.add_run(f"Q. {unit}: Answer the Following")
        unit_run.font.bold = True
        unit_run.font.size = Pt(12)
        unit_header.paragraph_format.space_after = Pt(12)

        # Create and configure table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Add header row with correct headers
        headers = ['Q. No.', 'Question', 'Difficulty Level', 'Blooms Level', 'CO', 'Marks']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(13)

        # Add questions
        subquestion = 0
        unit_questions = final_questions[final_questions['Unit_No'] == unit]
        for idx, question_row in unit_questions.iterrows():
            cells = table.add_row().cells

            # Set Q. No. (A, B, C, etc.)
            cells[0].text = chr(65 + subquestion % 26)
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            subquestion += 1

            # Add question content
            question_para = cells[1].paragraphs[0]
            convert_excel_to_word_with_images_and_equations(ws, question_row.name + 2, question_para)

            # Set other columns
            cells[2].text = str(question_row['Diff_Level'])
            cells[3].text = str(question_row['Blooms_Level'])
            cells[4].text = str(question_row['CO'])
            cells[5].text = str(int(question_row['Marks']))

            # Apply styling to data cells
            for i, cell in enumerate(cells):
                if i != 1:  # All except Question column
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)

        # Set column widths
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = template_table_widths[i]

        # Add spacing after table
        doc.add_paragraph()

    # Add end of paper
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_run = end_para.add_run("*** End of Paper ***")
    end_run.font.bold = True
    end_run.font.size = Pt(12)
    
    # Remove the template table
    if len(doc.tables) >= 3:
        doc.tables[2]._element.getparent().remove(doc.tables[2]._element)
    
    # Save document
    doc.save(word_file)
    print(f'Master Question paper has been successfully generated: {word_file}')